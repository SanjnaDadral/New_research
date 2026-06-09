import logging
import os
import re
import time
from typing import Dict, List
from pathlib import Path

from django.conf import settings

logger = logging.getLogger(__name__)

PDF_MAX_PAGES = int(os.getenv("PDF_MAX_PAGES", "30"))
PDF_PREFER_FAST = os.getenv("PDF_PREFER_FAST", "True").lower() == "true"
MIN_TEXT_CHARS_FOR_FAST_OK = 80


def _count_embedded_images_pypdf(reader, pages_to_read: int) -> int:
    n = 0
    for i in range(min(len(reader.pages), pages_to_read)):
        try:
            page = reader.pages[i]
            resources = page.get("/Resources")
            if not resources:
                continue

            xobjects = resources.get("/XObject")
            if not xobjects:
                continue

            for obj in xobjects.values():
                try:
                    if obj.get("/Subtype") == "/Image":
                        n += 1
                except:
                    continue
        except:
            continue
    return n


class PDFProcessor:
    def __init__(self):
        try:
            import fitz
            self._fitz_available = True
        except ImportError:
            self._fitz_available = False

        try:
            import pdfplumber
            self._pdfplumber_available = True
        except ImportError:
            self._pdfplumber_available = False

        try:
            import pypdf
            self._pypdf_available = True
        except ImportError:
            self._pypdf_available = False

    # =========================
    # MAIN ENTRY
    # =========================
    def extract_text(self, pdf_file) -> Dict[str, any]:
        try:
            if self._fitz_available:
                return self._extract_with_fitz(pdf_file)

            if PDF_PREFER_FAST and self._pypdf_available:
                fast = self._extract_with_pypdf(pdf_file)

                if fast["success"] and len(fast["text"].strip()) > MIN_TEXT_CHARS_FOR_FAST_OK:
                    return fast

                if self._pdfplumber_available:
                    return self._extract_with_pdfplumber(pdf_file)

                return fast

            if self._pdfplumber_available:
                return self._extract_with_pdfplumber(pdf_file)

            return self._extract_with_pypdf(pdf_file)

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return {"success": False, "error": str(e), "text": "", "pages": 0}

    # =========================
    # FITZ (PYMUPDF) - BEST
    # =========================
    def _extract_with_fitz(self, pdf_file) -> Dict[str, any]:
        import fitz
        try:
            pdf_file.seek(0)
            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            page_count = len(doc)
            pages_to_read = min(page_count, PDF_MAX_PAGES)
            
            text_parts = []
            extracted_images = []
            extracted_links = []
            
            image_dir = Path(settings.MEDIA_ROOT) / 'extracted' / 'images'
            image_dir.mkdir(parents=True, exist_ok=True)
            
            import time
            base_timestamp = int(time.time())
            
            for page_index in range(pages_to_read):
                page = doc[page_index]
                text_parts.append(page.get_text())
                
                # Extract Links
                links = page.get_links()
                for link in links:
                    if link.get("kind") == fitz.LINK_URI:
                        uri = link.get("uri")
                        if uri and uri not in extracted_links and uri.startswith("http"):
                            extracted_links.append(uri)
                
                # Extract Images
                if len(extracted_images) < 5:
                    images = page.get_images(full=True)
                    for img_index, img in enumerate(images):
                        if len(extracted_images) >= 5:
                            break
                        xref = img[0]
                        try:
                            base_image = doc.extract_image(xref)
                            if base_image:
                                image_bytes = base_image["image"]
                                image_ext = base_image["ext"]
                                if len(image_bytes) > 20000:  # Skip tiny icons / logos (< 20KB)
                                    filename = f"img_{base_timestamp}_{page_index}_{img_index}.{image_ext}"
                                    filepath = image_dir / filename
                                    with open(filepath, "wb") as f:
                                        f.write(image_bytes)
                                    # Normalize slashes for web URL
                                    media_path = f"{settings.MEDIA_URL}extracted/images/{filename}".replace('\\', '/')
                                    extracted_images.append(media_path)
                        except Exception as img_try_ex:
                            continue

            metadata = {}
            if doc.metadata:
                metadata = {
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                }

            return {
                "success": True,
                "text": "\n\n".join(text_parts),
                "pages": page_count,
                "pages_extracted": pages_to_read,
                "embedded_image_objects": len(extracted_images),
                "metadata": metadata,
                "extracted_images": extracted_images,
                "extracted_links": extracted_links[:10]
            }

        except Exception as e:
            logger.error(f"fitz error: {e}")
            return {"success": False, "error": str(e), "text": "", "pages": 0}

    # =========================
    # PDFPLUMBER (ACCURATE)
    # =========================
    def _extract_with_pdfplumber(self, pdf_file) -> Dict[str, any]:
        import pdfplumber

        try:
            pdf_file.seek(0)

            text_parts = []
            embedded_img = 0

            with pdfplumber.open(pdf_file) as pdf:
                page_count = len(pdf.pages)
                pages_to_read = min(page_count, PDF_MAX_PAGES)

                for page in pdf.pages[:pages_to_read]:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                    try:
                        embedded_img += len(page.images or [])
                    except:
                        pass

                metadata = self._extract_metadata_pdfplumber(pdf)

            return {
                "success": True,
                "text": "\n\n".join(text_parts),
                "pages": page_count,
                "pages_extracted": pages_to_read,
                "embedded_image_objects": embedded_img,
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"pdfplumber error: {e}")
            return {"success": False, "error": str(e), "text": "", "pages": 0}

    # =========================
    # PYPDF (FAST)
    # =========================
    def _extract_with_pypdf(self, pdf_file) -> Dict[str, any]:
        from pypdf import PdfReader

        try:
            pdf_file.seek(0)

            reader = PdfReader(pdf_file)
            page_count = len(reader.pages)
            pages_to_read = min(page_count, PDF_MAX_PAGES)

            text_parts = []

            for page in reader.pages[:pages_to_read]:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            metadata = {}
            if reader.metadata:
                metadata = {
                    "title": reader.metadata.get("/Title", ""),
                    "author": reader.metadata.get("/Author", ""),
                }

            return {
                "success": True,
                "text": "\n\n".join(text_parts),
                "pages": page_count,
                "pages_extracted": pages_to_read,
                "embedded_image_objects": _count_embedded_images_pypdf(reader, pages_to_read),
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"pypdf error: {e}")
            return {"success": False, "error": str(e), "text": "", "pages": 0}

    # =========================
    # METADATA
    # =========================
    def _extract_metadata_pdfplumber(self, pdf):
        try:
            return {
                "title": pdf.metadata.get("Title", ""),
                "author": pdf.metadata.get("Author", ""),
            }
        except:
            return {}

    # =========================
    # TITLE DETECTION
    # =========================
    def extract_title_from_pdf(self, text: str) -> str:
        lines = [l.strip() for l in text.split("\n") if l.strip()]

        for line in lines[:15]:
            if 10 < len(line) < 200:
                return line

        return "Untitled Document"


# =========================
# WORD (.docx/.doc) EXTRACTOR
# =========================
def extract_word_text(word_file) -> Dict[str, any]:
    """Extract text from a .docx or .doc file. Returns same shape as PDFProcessor.extract_text()."""
    filename = getattr(word_file, 'name', '').lower()

    # .docx — use python-docx
    if filename.endswith('.docx'):
        try:
            from docx import Document as DocxDocument
            word_file.seek(0)
            doc = DocxDocument(word_file)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            text = '\n\n'.join(paragraphs)
            # Extract core properties for metadata
            props = doc.core_properties
            metadata = {
                'title': props.title or '',
                'author': props.author or '',
            }
            return {
                'success': True,
                'text': text,
                'pages': len(doc.sections),
                'pages_extracted': len(doc.sections),
                'embedded_image_objects': 0,
                'metadata': metadata,
                'extracted_images': [],
                'extracted_links': [],
            }
        except Exception as e:
            logger.error(f"python-docx extraction error: {e}")
            return {'success': False, 'error': str(e), 'text': '', 'pages': 0}

    # .doc — try antiword via subprocess, fall back to raw text extraction
    if filename.endswith('.doc'):
        try:
            import subprocess, tempfile, os
            word_file.seek(0)
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
                tmp.write(word_file.read())
                tmp_path = tmp.name
            try:
                result = subprocess.run(
                    ['antiword', tmp_path],
                    capture_output=True, text=True, timeout=30
                )
                text = result.stdout.strip()
            finally:
                os.unlink(tmp_path)
            if text:
                return {
                    'success': True, 'text': text, 'pages': 1,
                    'pages_extracted': 1, 'embedded_image_objects': 0,
                    'metadata': {}, 'extracted_images': [], 'extracted_links': [],
                }
        except Exception:
            pass
        # Last resort: try opening as docx anyway (some .doc files are actually docx)
        try:
            from docx import Document as DocxDocument
            word_file.seek(0)
            doc = DocxDocument(word_file)
            text = '\n\n'.join(p.text for p in doc.paragraphs if p.text.strip())
            if text:
                return {
                    'success': True, 'text': text, 'pages': 1,
                    'pages_extracted': 1, 'embedded_image_objects': 0,
                    'metadata': {}, 'extracted_images': [], 'extracted_links': [],
                }
        except Exception:
            pass
        return {
            'success': False,
            'error': '.doc format has limited support. Please save as .docx and re-upload.',
            'text': '', 'pages': 0,
        }

    return {'success': False, 'error': 'Unsupported file type', 'text': '', 'pages': 0}


def is_word_file(filename: str) -> bool:
    return filename.lower().endswith(('.docx', '.doc'))


def is_pdf_file(filename: str) -> bool:
    return filename.lower().endswith('.pdf')


# =========================
# SINGLETON
# =========================
_pdf_processor = None


def get_pdf_processor():
    global _pdf_processor
    if _pdf_processor is None:
        _pdf_processor = PDFProcessor()
    return _pdf_processor